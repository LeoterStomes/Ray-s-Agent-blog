"""Export content to PDF / DOCX / TXT — save to uploads/export/ and return URL"""
import os, re, uuid, io
from datetime import datetime

EXPORT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "uploads", "export")
os.makedirs(EXPORT_DIR, exist_ok=True)


def _sanitize_filename(name: str) -> str:
    """移除文件名中的非法字符和空格"""
    name = re.sub(r'\s+', '-', name)  # 空格→连字符
    return re.sub(r'[\\/:*?"<>|]', '_', name)[:50]


def _html_to_text(html: str) -> str:
    """Strip HTML tags for plain text formatting"""
    t = re.sub(r'<br\s*/?>', '\n', html, flags=re.IGNORECASE)
    t = re.sub(r'</p>', '\n', t, flags=re.IGNORECASE)
    t = re.sub(r'</h[1-6]>', '\n\n', t, flags=re.IGNORECASE)
    t = re.sub(r'</li>', '\n', t, flags=re.IGNORECASE)
    t = re.sub(r'<[^>]+>', '', t)
    t = re.sub(r'\n{3,}', '\n\n', t)
    return t.strip()


def export_pdf(title: str, content: str) -> dict:
    """Generate PDF, return {url, filename}"""
    from fpdf import FPDF

    filename = f"{_sanitize_filename(title)}_{uuid.uuid4().hex[:6]}.pdf"
    filepath = os.path.join(EXPORT_DIR, filename)

    pdf = FPDF()
    pdf.add_page()
    # 注册中文字体（如果存在）
    font_paths = [
        os.path.join(os.path.dirname(__file__), "..", "fonts", "NotoSansSC-Regular.ttf"),
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/msyh.ttf",
    ]
    font_loaded = False
    for fp in font_paths:
        if os.path.exists(fp):
            pdf.add_font("CJK", "", fp, uni=True)
            pdf.add_font("CJK", "B", fp, uni=True)
            font_loaded = True
            break

    # Title
    if font_loaded:
        pdf.set_font("CJK", "B", 18)
    else:
        pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 14, title, ln=True)
    pdf.ln(6)

    # Date
    if font_loaded:
        pdf.set_font("CJK", "", 10)
    else:
        pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(128, 128, 128)
    pdf.cell(0, 8, f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}", ln=True)
    pdf.ln(6)
    pdf.set_text_color(0, 0, 0)

    # Content
    text = _html_to_text(content)
    if font_loaded:
        pdf.set_font("CJK", "", 11)
    else:
        pdf.set_font("Helvetica", "", 11)

    for line in text.split('\n'):
        line = line.strip()
        if not line:
            pdf.ln(4)
            continue
        if line.startswith('# ') or line.startswith('## ') or line.startswith('### '):
            pdf.ln(4)
            if font_loaded:
                pdf.set_font("CJK", "B", 13 if line.startswith('# ') else 12)
            pdf.cell(0, 8, line.lstrip('# '), ln=True)
            if font_loaded:
                pdf.set_font("CJK", "", 11)
            else:
                pdf.set_font("Helvetica", "", 11)
            pdf.ln(2)
        else:
            pdf.multi_cell(0, 6, line)
            pdf.ln(1)

    pdf.output(filepath)
    return {"filename": filename, "url": f"/uploads/export/{filename}", "type": "pdf"}


def export_docx(title: str, content: str) -> dict:
    """Generate DOCX, return {url, filename}"""
    from docx import Document
    from docx.shared import Pt, Inches

    filename = f"{_sanitize_filename(title)}_{uuid.uuid4().hex[:6]}.docx"
    filepath = os.path.join(EXPORT_DIR, filename)

    doc = Document()
    # Title
    h = doc.add_heading(title, level=0)
    # Date
    p = doc.add_paragraph(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = None  # gray

    # Content (basic HTML→DOCX conversion)
    text = content
    # Split by HTML headings
    sections = re.split(r'(<h[1-3][^>]*>.*?</h[1-3]>)', text, flags=re.IGNORECASE)
    for sec in sections:
        sec = sec.strip()
        if not sec:
            continue
        m = re.match(r'<h([1-3])[^>]*>(.*?)</h\1>', sec, re.IGNORECASE)
        if m:
            level = int(m.group(1))
            heading_text = re.sub(r'<[^>]+>', '', m.group(2)).strip()
            doc.add_heading(heading_text, level=min(level + 1, 3))
        else:
            clean = sec.replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')
            clean = re.sub(r'</p>', '\n\n', clean, re.IGNORECASE)
            clean = re.sub(r'</li>', '\n', clean, re.IGNORECASE)
            clean = re.sub(r'<[^>]+>', '', clean)
            clean = re.sub(r'\n{3,}', '\n\n', clean).strip()
            for para in clean.split('\n\n'):
                para = para.strip()
                if para:
                    p = doc.add_paragraph(para)

    doc.save(filepath)
    return {"filename": filename, "url": f"/uploads/export/{filename}", "type": "docx"}


def export_txt(title: str, content: str) -> dict:
    """Generate TXT, return {url, filename}"""
    filename = f"{_sanitize_filename(title)}_{uuid.uuid4().hex[:6]}.txt"
    filepath = os.path.join(EXPORT_DIR, filename)
    text = _html_to_text(content)
    text = f"{title}\n{'='*len(title)}\n导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n{text}"
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(text)
    return {"filename": filename, "url": f"/uploads/export/{filename}", "type": "txt"}
